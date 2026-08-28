import os
import json
import io
import streamlit as st
import google.generativeai as genai
import streamlit.components.v1 as components

st.set_page_config(page_title="뉴스로 배우는 영어회화", page_icon="🗞️", layout="centered")

try:
    api_key = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=api_key)
except Exception as e:
    st.error("⚠ Secrets에 GEMINI_API_KEY를 등록해 주세요.")
    st.stop()

desktop_path = os.path.join(os.path.expanduser("~"), "OneDrive", "Documents", "Desktop")
SAVE_DIR = os.path.join(desktop_path, "English_News_Notes")
os.makedirs(SAVE_DIR, exist_ok=True)

st.title("🗞️ 뉴스로 배우는 영어 회화 앱")
st.caption("London, Ontario 지역 뉴스 + 세계 뉴스로 매일 영어 토론 연습")

# 1. 뉴스 소스
st.subheader("1. 오늘의 뉴스로 바로 학습 (선택)")
news_titles = []
try:
    import feedparser
    local_feed = feedparser.parse("https://www.lfpress.com/feed/")
    world_feed = feedparser.parse("http://feeds.bbci.co.uk/news/world/rss.xml")
    local_titles = [f" [London, ON] {e.title}" for e in local_feed.entries[:3]]
    world_titles = [f" [World] {e.title}" for e in world_feed.entries[:3]]
    news_titles = local_titles + world_titles
    if news_titles:
        selected_news = st.selectbox("헤드라인을 선택하면 아래 입력창에 자동 입력됩니다", ["선택 안함"] + news_titles)
        if selected_news!= "선택 안함":
            clean_title = selected_news.replace(" [London, ON] ", "").replace(" [World] ", "")
            st.session_state['auto_keyword'] = clean_title
except:
    pass

# 2. 입력
col_level, col_input = st.columns([1, 2])
with col_level:
    level = st.selectbox("레벨", ["초급 - Slow & Simple", "중급 - Natural", "고급 - Debate / Discussion"], index=1)
with col_input:
    default_val = st.session_state.get('auto_keyword', '')
    news_input = st.text_input("뉴스 키워드 또는 기사 URL 입력", value=default_val, placeholder="예: London downtown construction")

url_input = st.text_input("선택: 뉴스 기사 원문 URL (있으면 더 정확함)", placeholder="https://lfpress.com/...")

if "news_result" not in st.session_state:
    st.session_state.news_result = ""
if "current_news" not in st.session_state:
    st.session_state.current_news = ""

if st.button("📰 회화 자료 생성하기") and (news_input or url_input):
    with st.spinner("AI가 뉴스 기반 회화 자료를 만들고 있습니다..."):
        try:
            # --- 수정된 부분: 3.6 버전처럼 안정적인 모델 이름으로 복구 ---
            model = genai.GenerativeModel(
                model_name="gemini-3.6-flash", # 1.5-flash가 v1beta에서 404 에러나서 2.0-flash로 변경 (기존 3.6 호환)
                system_instruction=f"You are an expert English conversation coach for Korean learners in London, ON. Level: {level}. Provide 1. Summary EN+KO, 2. Key Expressions 5, 3. Role-Play Dialogue A:B with EN+KO.",
                generation_config={"response_mime_type": "text/plain"}
            )
            prompt = f"News: {news_input} URL hint: {url_input} Level: {level}"
            response = model.generate_content(prompt)
            st.session_state.news_result = response.text
            st.session_state.current_news = news_input or url_input
        except Exception as e:
            st.error(f"오류: {e}")

def create_docx_bytes(title, content):
    from docx import Document
    doc = Document()
    doc.add_heading(f'뉴스 영어회화 노트: {title}', 1)
    clean = content.replace('**','').replace('###','').replace('##','').replace('---','')
    for line in clean.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

if st.session_state.news_result:
    st.markdown(f"## ✨ {st.session_state.current_news}")
    st.markdown(st.session_state.news_result)
    st.markdown("---")

    safe_text_json = json.dumps(st.session_state.news_result)
    tts_html = f"""
    <div style="display: flex; gap: 10px; flex-wrap: wrap; margin-bottom:10px;">
        <button onclick="speakDialogueToneShift()" style="background-color: #4CAF50; color: white; padding: 10px 15px; border: none; border-radius: 5px; cursor: pointer; font-weight: bold;">👥 토론 대화 입체 톤으로 듣기</button>
        <button onclick="stopSpeech()" style="background-color: #f44336; color: white; padding: 10px 15px; border: none; border-radius: 5px; cursor: pointer; font-weight: bold;">⏹ 중지</button>
    </div>
    <script>
    const fullMarkdownText = {safe_text_json};
    function extractEnglishLines(s){{ let l=s.split('\\n'), e=[]; for(let line of l){{ let k=/[가-힣]/.test(line), en=/[a-zA-Z]/.test(line); if(en&&!k){{ let c=line.replace(/^[#*\\-\\d\\.\\s]+/, '').replace(/^[A-Za-z]\\s*:/, '').trim(); if(c.length>0) e.push(c); }} }} return e; }}
    function speakDialogueToneShift(){{ window.speechSynthesis.cancel(); let d=fullMarkdownText, di=d.indexOf('3.'), di2=d.indexOf('Role-Play'); let start = (di2!=-1)?di2:di; if(start!=-1) d=d.substring(start); let lines=extractEnglishLines(d); let i=0; function nxt(){{ if(i>=lines.length) return; let u=new SpeechSynthesisUtterance(lines[i]); u.lang='en-US'; u.rate=0.9; u.pitch=(i%2==0)?1.2:0.8; u.onend=function(){{i++; nxt();}}; window.speechSynthesis.speak(u); }} nxt(); }}
    function stopSpeech(){{ window.speechSynthesis.cancel(); }}
    </script>
    """
    components.html(tts_html, height=90)

    file_name_docx = f"{st.session_state.current_news[:20].lower().replace(' ','_')}_news.docx"
    file_name_md = f"{st.session_state.current_news[:20].lower().replace(' ','_')}_news.md"
    note_md = f"# 뉴스 영어회화 노트: {st.session_state.current_news}\n\n" + st.session_state.news_result

    col1, col2 = st.columns(2)
    with col1:
        docx_bytes = create_docx_bytes(st.session_state.current_news, st.session_state.news_result)
        st.download_button(label="📥 기기에 다운로드", data=docx_bytes, file_name=file_name_docx, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with col2:
        if st.button("☁ 클라우드에 저장"):
            path = os.path.join(SAVE_DIR, file_name_md)
            with open(path, "w", encoding="utf-8") as f:
                f.write(note_md)
            st.success(f"저장 완료!")

# --- 사이드바: 이전처럼 프린트 기능 복구 ---
st.sidebar.markdown("---")
st.sidebar.subheader("📁 저장된 뉴스 노트")
if os.path.exists(SAVE_DIR):
    files = [f for f in os.listdir(SAVE_DIR) if f.endswith('.md')]
    if files:
        sel = st.sidebar.selectbox("파일 선택", files)
        if sel:
            with open(os.path.join(SAVE_DIR, sel), "r", encoding="utf-8") as f:
                c = f.read()
            st.sidebar.markdown(c[:2000])

            print_html = f"""
            <button onclick="printNote()" style="width:100%; background-color:#607D8B; color:white; padding:8px; border:none; border-radius:5px; cursor:pointer;">🖨️ 이 노트 프린트 하기</button>
            <div id="print-content" style="display:none;">{c.replace(chr(10), '<br>')}</div>
            <script>
            function printNote(){{
                var w = window.open('', '', 'height=600,width=800');
                w.document.write('<html><head><title>News Note</title></head><body>');
                w.document.write(document.getElementById('print-content').innerHTML);
                w.document.write('</body></html>');
                w.document.close();
                w.print();
            }}
            </script>
            """
            with st.sidebar:
                components.html(print_html, height=80)
    else:
        st.sidebar.caption("저장된 노트가 없습니다.")
