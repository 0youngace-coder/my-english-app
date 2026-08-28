import os
import json
import io
import streamlit as st
import google.generativeai as genai
import streamlit.components.v1 as components

st.set_page_config(page_title="AI 영단어 학습 앱", page_icon="📚", layout="centered")

try:
    api_key = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=api_key)
except Exception as e:
    st.error("⚠ API 키 설정이 필요합니다. Streamlit Cloud Secrets에 GEMINI_API_KEY를 등록해 주세요.")
    st.stop()

desktop_path = os.path.join(os.path.expanduser("~"), "OneDrive", "Documents", "Desktop")
SAVE_DIR = os.path.join(desktop_path, "English_Notes")
if not os.path.exists(SAVE_DIR):
    os.makedirs(SAVE_DIR, exist_ok=True)

st.title("📚 AI 영단어 학습 앱 (스마트폰 최적화)")
st.write("단어를 검색한 뒤, 대화문 또는 소설문 음성 버튼을 눌러 들어보세요.")

word_input = st.text_input("영단어 입력 (예: irrevocable):", "")
if "result_text" not in st.session_state:
    st.session_state.result_text = ""
if "current_word" not in st.session_state:
    st.session_state.current_word = ""

if st.button("생성하기") and word_input:
    with st.spinner("AI가 멋진 예문과 대화문을 작성 중입니다..."):
        try:
            model = genai.GenerativeModel(
                model_name="gemini-3.6-flash",
                system_instruction="""You are an expert English teacher and writer.
                When given an English word, provide the output strictly in this format with clear headings:
                1. Meaning in Korean
                2. Dialogue (Provide a natural dialogue between native speakers using the word, using ONLY 'A:' and 'B:' for speakers, including English lines and Korean translations)
                3. Novel Passage (Provide a sophisticated literary/novel passage using the word, including English and Korean translation)
                Return the response in clean Markdown format.""",
                generation_config={"response_mime_type": "text/plain"}
            )
            response = model.generate_content(f"Word: {word_input}")
            st.session_state.result_text = response.text
            st.session_state.current_word = word_input.strip()
        except Exception as e:
            st.error(f"오류가 발생했습니다: {e}")

# --- 결과가 있을 때만 표시 (버그 수정 포인트) ---
if st.session_state.result_text:
    st.markdown(f"## ✨ {st.session_state.current_word}")
    st.markdown(st.session_state.result_text)
    st.markdown("---")

    safe_text_json = json.dumps(st.session_state.result_text)
    tts_html = f"""
    <div style="margin-bottom: 10px; display: flex; gap: 10px; flex-wrap: wrap;">
        <button onclick="speakDialogueToneShift()" style="background-color: #4CAF50; color: white; padding: 10px 15px; border: none; border-radius: 5px; cursor: pointer; font-size: 15px; font-weight: bold;">👥 대화문 입체 톤으로 듣기</button>
        <button onclick="speakNovel()" style="background-color: #2196F3; color: white; padding: 10px 15px; border: none; border-radius: 5px; cursor: pointer; font-size: 15px; font-weight: bold;">📖 영어 소설문만 듣기</button>
        <button onclick="stopSpeech()" style="background-color: #f44336; color: white; padding: 10px 15px; border: none; border-radius: 5px; cursor: pointer; font-size: 15px; font-weight: bold;">⏹ 중지</button>
    </div>
    <script>
    const fullMarkdownText = {safe_text_json};
    function extractEnglishLines(textSection) {{
        let lines = textSection.split('\\n'); let englishLines = [];
        for (let line of lines) {{
            let hasKorean = /[가-힣]/.test(line); let hasEnglish = /[a-zA-Z]/.test(line);
            if (hasEnglish &&!hasKorean) {{
                let cleaned = line.replace(/^[#*\\-\\d\\.\\s]+/, '').replace(/^[A-Za-z]\\s*:/, '').trim();
                if (cleaned.length > 0) englishLines.push(cleaned);
            }}
        }}
        return englishLines;
    }}
    function speakDialogueToneShift() {{
        if ('speechSynthesis' in window) {{
            window.speechSynthesis.cancel();
            let dialoguePart = fullMarkdownText; let dIndex = fullMarkdownText.indexOf('2.'); let nIndex = fullMarkdownText.indexOf('3.');
            if (dIndex!== -1 && nIndex!== -1) dialoguePart = fullMarkdownText.substring(dIndex, nIndex);
            else if (dIndex!== -1) dialoguePart = fullMarkdownText.substring(dIndex);
            let lines = extractEnglishLines(dialoguePart); if (lines.length === 0) return; let index = 0;
            function speakNextLine() {{
                if (index >= lines.length) return;
                let utterance = new SpeechSynthesisUtterance(lines[index]); utterance.lang = 'en-US'; utterance.rate = 0.9; utterance.pitch = (index % 2 === 0)? 1.2 : 0.8;
                utterance.onend = function() {{ index++; speakNextLine(); }}; window.speechSynthesis.speak(utterance);
            }}
            speakNextLine();
        }}
    }}
    function speakNovel() {{
        if ('speechSynthesis' in window) {{
            window.speechSynthesis.cancel(); let novelPart = ""; let nIndex = fullMarkdownText.indexOf('3.');
            if (nIndex === -1) nIndex = fullMarkdownText.indexOf('3'); if (nIndex!== -1) novelPart = fullMarkdownText.substring(nIndex); else novelPart = fullMarkdownText;
            let lines = extractEnglishLines(novelPart); let textToRead = lines.join('. '); if (!textToRead || textToRead.trim().length < 2) return;
            var utterance = new SpeechSynthesisUtterance(textToRead); utterance.lang = 'en-US'; utterance.rate = 0.85; window.speechSynthesis.speak(utterance);
        }}
    }}
    function stopSpeech() {{ if ('speechSynthesis' in window) window.speechSynthesis.cancel(); }}
    </script>
    """
    components.html(tts_html, height=160)

    # --- 다운로드 기능 (MD + Word + PDF) ---
    target_word = st.session_state.current_word or "note"
    base_name = target_word.lower()
    note_content_md = f"# 영단어 학습 노트: {target_word}\n\n" + st.session_state.result_text

    def create_docx_bytes(word, content):
        try:
            from docx import Document
            doc = Document()
            doc.add_heading(f'영단어 학습 노트: {word}', 1)
            doc.add_paragraph(content)
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception as e:
            return None

    def create_pdf_bytes(word, content):
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            buf = io.BytesIO()
            doc = SimpleDocTemplate(buf, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            story.append(Paragraph(f"English Note: {word}", styles['Title']))
            story.append(Spacer(1, 12))
            # Markdown 기호 제거해서 깔끔하게
            clean_content = content.replace('**','').replace('###','').replace('##','')
            for para in clean_content.split('\n'):
                if para.strip():
                    story.append(Paragraph(para, styles['Normal']))
                    story.append(Spacer(1, 6))
            doc.build(story)
            return buf.getvalue()
        except Exception as e:
            return None

    st.markdown("#### 💾 저장하기")
    c1, c2, c3 = st.columns(3)
    with c1:
        st.download_button("📝 Word로 저장", data=create_docx_bytes(target_word, st.session_state.result_text) or note_content_md, file_name=f"{base_name}_note.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with c2:
        pdf_bytes = create_pdf_bytes(target_word, st.session_state.result_text)
        if pdf_bytes:
            st.download_button("📄 PDF로 저장", data=pdf_bytes, file_name=f"{base_name}_note.pdf", mime="application/pdf")
        else:
            st.caption("PDF 라이브러리 필요")
    with c3:
        st.download_button("📥 MD로 저장", data=note_content_md, file_name=f"{base_name}_note.md", mime="text/markdown")

# --- 사이드바 (원본 인쇄 기능 유지) ---
st.sidebar.markdown("---")
st.sidebar.subheader("📁 저장된 노트 보관함")
st.sidebar.write(f"폴더 위치:\n`{SAVE_DIR}`")

if os.path.exists(SAVE_DIR):
    saved_files = [f for f in os.listdir(SAVE_DIR) if f.endswith('.md')]
    if saved_files:
        st.sidebar.write(f"총 저장된 파일: {len(saved_files)}개")
        selected_file = st.sidebar.selectbox("파일 선택하여 보기", saved_files)
        if selected_file:
            with open(os.path.join(SAVE_DIR, selected_file), "r", encoding="utf-8") as f:
                content = f.read()
            st.sidebar.markdown("---")
            st.sidebar.markdown(content)
            print_html = f"""
                <div style="font-family: Arial, sans-serif; padding: 20px; line-height: 1.6; max-width: 800px; margin: 0 auto;">
                    <h3 style="color: #333;">📖 학습 노트: {selected_file}</h3>
                    <hr style="border: 1px solid #ddd; margin-bottom: 20px;">
                    <div style="white-space: pre-wrap; font-size: 16px; color: #111;">{content}</div>
                </div>
                <script>window.print();</script>
            """
            if st.sidebar.button("🖨 이 노트 프린트하기"):
                components.html(print_html, height=400)
